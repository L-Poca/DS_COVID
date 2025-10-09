# word_generator.py
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime


def create_report_folder(folder_name="rapports"):
    """Crée un dossier 'rapports' s'il n'existe pas."""
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
    return folder_name


def generate_word_report(models_results, output_folder="rapports"):
    """
    Génère un rapport Word avec :
    - Titre
    - Matrices de confusion (images)
    - Rapports de classification (texte)

    models_results : liste de dicts
        [
            {
                'name': 'Random Forest',
                'report': "precision    recall  f1-score ...",
                'confusion_matrix_path': 'cm_rf.png'
            },
            ...
        ]
    """
    # Créer le dossier de sortie
    report_folder = create_report_folder(output_folder)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    doc_path = os.path.join(
        report_folder,
        f"rapport_classification_{timestamp}.docx"
        )

    # Créer le document Word
    doc = Document()
    doc.add_heading('Rapport de Classification - Modèles ML', 0)
    doc.add_paragraph(
        f"Généré le : {datetime.now().strftime('%d/%m/%Y à %H:%M:%S')}"
        )
    doc.add_paragraph(
        "Modèles évalués sur le dataset COVID-19 Radiography (4 classes)."
        )

    for result in models_results:
        doc.add_heading(f"Modèle : {result['name']}", level=1)

        # Ajouter le rapport de classification en tant que texte
        doc.add_heading("Rapport de classification", level=2)
        doc.add_paragraph(result['report'])

        # Ajouter la matrice de confusion
        if 'confusion_matrix_path' in result and os.path.exists(
            result['confusion_matrix_path']
             ):
            doc.add_heading("Matrice de confusion", level=2)
            doc.add_picture(result['confusion_matrix_path'], width=Inches(4.5))
        else:
            doc.add_paragraph(
                "[Matrice de confusion non disponible]",
                style='Italic'
                )

        doc.add_page_break()  # Optionnel : une page par modèle

    doc.save(doc_path)
    print(f"\n✅ Rapport Word généré : {doc_path}")
    return doc_path
