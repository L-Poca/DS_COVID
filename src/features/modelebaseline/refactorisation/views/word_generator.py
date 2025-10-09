# word_generator.py
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime
try:
    from docx2pdf import convert  # conversion Word → PDF
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    print("[WARNING] docx2pdf non installé. Les rapports PDF ne seront pas générés.")


def create_report_folder(folder_name="rapports"):
    """Crée un dossier 'rapports' s'il n'existe pas."""
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)
    return folder_name


def generate_word_report(models_results, output_folder="rapports", generate_pdf=True):
    """
    Génère un rapport Word et optionnellement PDF avec :
    - Titre
    - Matrices de confusion (images)
    - Rapports de classification (texte)

    models_results : liste de dicts
        [
            {
                'name': 'Random Forest',
                'report': "precision    recall  f1-score ...",
                'confusion_matrix_path': 'cm_rf.png'
            },
            ...
        ]
    """
    # Créer le dossier de sortie
    report_folder = create_report_folder(output_folder)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = os.path.join(report_folder, f"rapport_classification_{timestamp}.docx")
    pdf_path = os.path.join(report_folder, f"rapport_classification_{timestamp}.pdf")

    # Créer le document Word
    doc = Document()
    doc.add_heading('Rapport de Classification - Modèles ML', 0)
    doc.add_paragraph(f"Généré le : {datetime.now().strftime('%d/%m/%Y à %H:%M:%S')}")
    doc.add_paragraph("Modèles évalués sur le dataset COVID-19 Radiography (4 classes).")

    for result in models_results:
        doc.add_heading(f"Modèle : {result['name']}", level=1)

        # Rapport de classification
        doc.add_heading("Rapport de classification", level=2)
        doc.add_paragraph(result['report'])

        # Matrice de confusion
        if 'confusion_matrix_path' in result and os.path.exists(result['confusion_matrix_path']):
            doc.add_heading("Matrice de confusion", level=2)
            doc.add_picture(result['confusion_matrix_path'], width=Inches(4.5))
        else:
            doc.add_paragraph("[Matrice de confusion non disponible]", style='Italic')

        doc.add_page_break()

    # Sauvegarde Word
    doc.save(docx_path)
    print(f"\n✅ Rapport Word généré : {docx_path}")

    # Conversion PDF si possible
    if generate_pdf and DOCX2PDF_AVAILABLE:
        convert(docx_path, pdf_path)
        print(f"✅ Rapport PDF généré : {pdf_path}")
    elif generate_pdf:
        print("[WARNING] docx2pdf non disponible → PDF non généré.")

    return docx_path, pdf_path if generate_pdf and DOCX2PDF_AVAILABLE else None
